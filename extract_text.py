"""
extract_text.py
-----------------
Pulls plain text out of a PDF or DOCX file.

This module supports both normal PDFs and scanned/image-based PDFs by
falling back to OCR when the raw PDF text is blank or very sparse.

For handwriting or cursive notes, Tesseract alone is limited. This version
adds image preprocessing and a better fallback engine when EasyOCR is
available.
"""

import os
import tempfile

import pdfplumber
import docx  # python-docx

try:
    import pytesseract
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None

try:
    from pdf2image import convert_from_path
except Exception:  # pragma: no cover - optional dependency
    convert_from_path = None

try:
    from PIL import Image, ImageFilter, ImageOps
except Exception:  # pragma: no cover - optional dependency
    Image = None
    ImageFilter = None
    ImageOps = None

try:
    import easyocr
except Exception:  # pragma: no cover - optional dependency
    easyocr = None

try:
    from google.cloud import vision
except Exception:  # pragma: no cover - optional dependency
    vision = None


def _configure_tesseract():
    if pytesseract is None or os.name != "nt":
        return

    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            break


def _preprocess_image(image):
    if Image is None:
        return image

    img = image.convert("L")
    img = ImageOps.autocontrast(img)
    img = img.resize((img.width * 3, img.height * 3))
    img = img.filter(ImageFilter.SHARPEN)

    # more aggressive thresholding for handwritten/scanned notes
    img = img.point(lambda px: 255 if px > 200 else 0)
    return img


def _ocr_image_with_tesseract(image) -> str:
    if pytesseract is None:
        return ""

    _configure_tesseract()
    processed = _preprocess_image(image)
    configs = [
        "--psm 6 --oem 3",
        "--psm 11 --oem 3",
        "--psm 4 --oem 3",
        "--psm 3 --oem 3",
    ]

    for config in configs:
        try:
            text = pytesseract.image_to_string(processed, config=config)
            cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
            if cleaned:
                return cleaned
        except Exception:
            continue
    return ""


def _ocr_image_with_easyocr(image) -> str:
    if easyocr is None:
        return ""

    try:
        reader = easyocr.Reader(["en"], gpu=False)
        results = reader.readtext(image, detail=0, paragraph=True)
        text = "\n".join(part.strip() for part in results if part and part.strip())
        if text:
            return text
    except Exception:
        pass
    return ""


def _ocr_image_with_google_vision(image_path: str) -> str:
    if vision is None:
        return ""

    if not os.getenv("GOOGLE_APPLICATION_CREDENTIALS"):
        return ""

    try:
        client = vision.ImageAnnotatorClient()
        with open(image_path, "rb") as f:
            content = f.read()
        image = vision.Image(content=content)

        # document_text_detection is more reliable for handwritten and scanned notes
        for method_name in ("document_text_detection", "text_detection"):
            method = getattr(client, method_name, None)
            if method is None:
                continue

            response = method(image=image)
            if method_name == "document_text_detection":
                text = getattr(response, "full_text_annotation", None)
                if text is not None and text.text:
                    return text.text.strip()

            texts = getattr(response, "text_annotations", None) or []
            if texts:
                return texts[0].description.strip()

        return ""
    except Exception:
        return ""


def _ocr_pdf(path: str) -> str:
    if convert_from_path is None:
        raise ValueError(
            f"OCR support is not installed for {path}. Install pdf2image and Tesseract, "
            "or configure Google Vision credentials."
        )

    images = convert_from_path(path, dpi=400)
    pages = []
    for image in images:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            image.save(tmp.name)
            tmp_path = tmp.name

        try:
            text = _ocr_image_with_google_vision(tmp_path)
            if not text:
                text = _ocr_image_with_tesseract(image)
            if not text:
                text = _ocr_image_with_easyocr(image)
            if not text:
                text = _ocr_image_with_google_vision(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        if text:
            pages.append(text.strip())

    text = "\n".join(pages).strip()
    if not text:
        raise ValueError(
            f"OCR could not extract any readable text from {path}. "
            "This usually means the scan is too faint, skewed, or handwritten and needs cleaner image input."
        )
    return text


def extract_pdf(path: str) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    text = "\n".join(text_parts).strip()

    if len(text) < 50:
        return _ocr_pdf(path)
    return text


def extract_docx(path: str) -> str:
    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of any tables, since docx tables aren't part of
    # document.paragraphs and would otherwise be silently skipped.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError(f"No extractable text found in {path}.")
    return text


def extract_text(path: str) -> str:
    """Single entry point - detects file type from the extension and routes
    to the right extractor."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext == ".docx":
        return extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .pdf and .docx are supported right now.")
