"""
save_output.py
----------------
Writes the simplified summary out as a clean .docx file.
"""

import docx
from docx.shared import Pt, RGBColor


def save_summary_docx(summary_text: str, output_path: str, source_filename: str):
    document = docx.Document()

    title = document.add_heading("Simplified Summary", level=1)

    subtitle = document.add_paragraph()
    run = subtitle.add_run(f"Source document: {source_filename}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_paragraph()  # spacer

    # summarize_text() returns a "Summary: ... \n\nKey Points:\n- ..." block
    # (or, for single-chunk documents, plain prose + bullets) - split it into
    # paragraphs/bullets so the docx looks structured rather than one wall of text.
    for line in summary_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("-") or line.startswith("*"):
            document.add_paragraph(line.lstrip("-* ").strip(), style="List Bullet")
        elif line.endswith(":") and len(line) < 40:
            document.add_heading(line.rstrip(":"), level=2)
        else:
            document.add_paragraph(line)

    document.save(output_path)
